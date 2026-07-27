"""Shared one-page resume/cover-letter builder for the job application pipeline.

Uses python-docx. Styled for ATS parseability per references/ats-optimization.md:
standard fonts, no tables-for-layout, no headers/footers for contact info, no icons.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x2C, 0x3E, 0x50)


def _set_margins(doc, top=0.5, bottom=0.5, left=0.6, right=0.6):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _no_space(paragraph, before=0, after=0, line=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def new_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.font.color.rgb = INK
    _set_margins(doc)
    return doc


def add_name_header(doc, name, contact_line):
    p = doc.add_paragraph()
    _no_space(p, after=2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    _no_space(p2, after=8)
    run2 = p2.add_run(contact_line)
    run2.font.size = Pt(9.5)


def add_title_line(doc, title_text):
    p = doc.add_paragraph()
    _no_space(p, after=8)
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT


def _add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C3E50")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    _no_space(p, before=6, after=3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = ACCENT
    _add_rule(p)


def add_summary(doc, text):
    p = doc.add_paragraph()
    _no_space(p, after=4)
    run = p.add_run(text)
    run.font.size = Pt(10)


def add_skills_line(doc, category, items):
    p = doc.add_paragraph()
    _no_space(p, after=2)
    r1 = p.add_run(f"{category}: ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r2 = p.add_run(items)
    r2.font.size = Pt(9.5)


def add_job_header(doc, title, company, dates):
    p = doc.add_paragraph()
    _no_space(p, before=5, after=1)
    r1 = p.add_run(f"{title} — {company}")
    r1.bold = True
    r1.font.size = Pt(10)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.8), alignment=2)
    p.add_run("\t")
    r2 = p.add_run(dates)
    r2.italic = True
    r2.font.size = Pt(9)


def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    _no_space(p, after=1)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.14)
    run = p.add_run(text)
    run.font.size = Pt(size)


def add_plain_line(doc, text, size=9.5, bold=False, italic=False):
    p = doc.add_paragraph()
    _no_space(p, after=2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


# ---- Cover letter helpers ----

def add_cover_date(doc, date_str):
    p = doc.add_paragraph()
    _no_space(p, after=8)
    run = p.add_run(date_str)
    run.font.size = Pt(10)


def add_cover_paragraph(doc, text, after=8):
    p = doc.add_paragraph()
    _no_space(p, after=after, line=1.15)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def save(doc, path):
    doc.save(path)
